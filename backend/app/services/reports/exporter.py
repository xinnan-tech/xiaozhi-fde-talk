"""报告导出：md / html / word。pdf 后加（需 weasyprint/pandoc）。

报告本就是 Markdown，导出是格式转换。
"""
from __future__ import annotations

import io
import re

import bleach
import markdown

from app.core.i18n import Keys
from app.core.i18n.errors import I18nError

FORMATS = ("md", "html", "word")


# Word docx 字体映射：按 llm.output_language 选 ascii/hAnsi（西文）+ eastAsia（中日韩）。
#
# 来源：用户反馈下载的 Word 字体是「MS 明朝」（Word 默认 eastAsiaTheme=minorEastAsia
# 解析结果）→ 中文场景应该是「宋体」，英文场景应该是「Times New Roman」。
#
# 选型理由：
# - zh_cn / zh_tw：宋体——Windows / macOS / Linux 主流 Office 都自带，访谈报告最常见选择。
# - en：Times New Roman——全系统预装、serif、正式/学术风。
# - 其他拉丁语种（vi/fr/de/es/ru）同 Times New Roman——拉丁字符集共用。
# - ja / ko：MS Mincho（ja 默认）/ Malgun Gothic（ko 默认）暂不覆盖，沿用 Times New Roman。
#   后续如需「日文报告 → 明朝源」可在此 dict 加分支。
# - 未知 lang：默认 Times New Roman + 宋体 兜底——避免 LLM 未知语种落到 MS 明朝。
_FONT_BY_LANG: dict[str, tuple[str, str]] = {
    "zh_cn": ("宋体", "宋体"),       # (ascii=宋体也兜底中文, eastAsia=宋体)
    "zh_tw": ("宋体", "宋体"),       # 繁中传统上用「細明體」，但宋体也能渲染简繁
    "en":    ("Times New Roman", "宋体"),
    "vi":    ("Times New Roman", "宋体"),
    "ru":    ("Times New Roman", "宋体"),
    "ko":    ("Times New Roman", "宋体"),
    "ja":    ("Times New Roman", "宋体"),
    "fr":    ("Times New Roman", "宋体"),
    "de":    ("Times New Roman", "宋体"),
    "es":    ("Times New Roman", "宋体"),
}


def _fonts_for(language: str) -> tuple[str, str]:
    """按 llm.output_language 取 (ascii_font, east_asia_font)。未知 lang 走 en 兜底。"""
    return _FONT_BY_LANG.get((language or "").lower(), _FONT_BY_LANG["en"])


class ReportFormatNotImplementedError(I18nError):
    """指定 format 暂未实现（如 pdf）。由 FastAPI I18nError handler 转结构化 501。"""
    def __init__(self, *, fmt: str):
        super().__init__(Keys.REPORT_FORMAT_NOT_IMPLEMENTED, http_status=501, fmt=fmt)
        # Backward compat: existing consumers (routes/reports.py + tests) read
        # `.fmt` directly; preserve it as a top-level attribute alongside `params['fmt']`.
        self.fmt = fmt


# Backward-compatible alias: `app/transport/http/routes/reports.py` and
# `tests/unit/test_report_export_pdf_501.py` still import this name. It is the
# SAME class as ReportFormatNotImplementedError so any `except` continues to work.
FormatNotImplementedError = ReportFormatNotImplementedError


_ALLOWED_TAGS = [
    "p", "br", "strong", "em", "ul", "ol", "li",
    "table", "thead", "tbody", "tr", "th", "td",
    "pre", "code", "h1", "h2", "h3", "h4", "blockquote", "a", "hr",
]
_ALLOWED_ATTRS = {"a": ["href"], "code": ["class"], "pre": ["class"]}


def _sanitize(html: str) -> str:
    return bleach.clean(html, tags=_ALLOWED_TAGS, attributes=_ALLOWED_ATTRS, strip=True)


def export(md: str, fmt: str, language: str = "en") -> tuple[bytes, str]:
    """Markdown → 指定格式，返回 (data, media_type)。

    language：仅 word 格式用——按 llm.output_language 选 ascii/eastAsia 字体。
    md / html 不读 language（html 用 system-ui 兜底，md 是纯文本）。
    """
    fmt = (fmt or "md").lower()
    if fmt == "md":
        return md.encode("utf-8"), "text/markdown; charset=utf-8"
    if fmt == "html":
        body = markdown.markdown(md, extensions=["tables", "fenced_code"])
        body = _sanitize(body)
        page = (
            "<!DOCTYPE html><html><head><meta charset='utf-8'>"
            "<style>body{font-family:system-ui,sans-serif;max-width:760px;margin:40px auto;padding:0 16px;}"
            "table{border-collapse:collapse;}td,th{border:1px solid #ccc;padding:4px 8px;}</style>"
            "</head><body>" + body + "</body></html>"
        )
        return page.encode("utf-8"), "text/html; charset=utf-8"
    if fmt == "word":
        return _to_docx(md, language), (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    if fmt == "pdf":
        raise ReportFormatNotImplementedError(fmt=fmt)
    raise I18nError(
        Keys.HTTP_REPORT_FORMAT_UNSUPPORTED, http_status=400,
        fmt=fmt, supported=FORMATS,
    )


def _apply_docx_fonts(doc, ascii_font: str, east_asia_font: str) -> None:
    """把 docx 里所有 rFonts 的 theme refs 替换成 named fonts。

    Word 默认 docDefaults/rPrDefault 用 asciiTheme="minorHAnsi" + eastAsiaTheme="minorEastAsia"，
    这俩 theme 默认解析为「Calibri Light」+「MS 明朝」——前者问题不大，后者下载用户一看就懵。
    修法：递归遍历 styles.xml 里所有 <w:rFonts>，清掉 asciiTheme/hAnsiTheme/eastAsiaTheme/cstheme
    四个 theme ref 属性，填入 named fonts（ascii/hAnsi/cs=latin, eastAsia=CJK）。
    """
    from docx.oxml.ns import qn

    styles_elem = doc.styles.element
    for rFonts in styles_elem.iter(qn("w:rFonts")):
        for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
            if rFonts.get(qn(attr)) is not None:
                del rFonts.attrib[qn(attr)]
        rFonts.set(qn("w:ascii"), ascii_font)
        rFonts.set(qn("w:hAnsi"), ascii_font)
        rFonts.set(qn("w:eastAsia"), east_asia_font)
        rFonts.set(qn("w:cs"), ascii_font)


def _to_docx(md: str, language: str = "en") -> bytes:
    """简易 Markdown → docx（标题/段落/列表/引用）。

    language 决定字体：
    - zh_cn/zh_tw → ascii/hAnsi=宋体, eastAsia=宋体（中文场景整篇中文，宋体即可）
    - en/vi/ru/... → ascii/hAnsi=Times New Roman, eastAsia=宋体（CJK 字符仍走宋体兜底）

    行内 `<br>` 会被翻译成 docx 真换行（add_break），保证 markdown-table 单元格里
    通过 `_escape_table_cell` 嵌入的 `<br>` 在 Word 里也是可见换行而非字面量。
    Markdown 表格 (`| ... | ... |`) 在 Word 渲染仍未支持——单元格会被当成普通段落，
    完整表格支持见后续 issue。
    """
    from docx import Document
    from docx.enum.text import WD_BREAK

    ascii_font, east_asia_font = _fonts_for(language)
    doc = Document()
    _apply_docx_fonts(doc, ascii_font, east_asia_font)
    for line in md.splitlines():
        s = line.rstrip()
        if not s:
            continue
        if s.startswith("### "):
            doc.add_heading(s[4:].strip(), level=3)
        elif s.startswith("## "):
            doc.add_heading(s[3:].strip(), level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:].strip(), level=1)
        elif s.startswith("> "):
            doc.add_paragraph(s[2:].strip(), style="Intense Quote" if "Intense Quote" in [st.name for st in doc.styles] else None)
        elif s.startswith(("- ", "* ")):
            doc.add_paragraph(s[2:].strip(), style="List Bullet")
        else:
            _add_paragraph_with_breaks(doc, s)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_paragraph_with_breaks(doc, text: str) -> None:
    """添加段落：把 `<br>` / `<br/>` / `<br />` 切成多段 run，用真换行连接。

    其余文本原样放入同一段。Markdown 表格行（`| a | b |`）也会走这里——后续
    issue 里给表格单独走 docx table 渲染时再覆盖。
    """
    from docx.enum.text import WD_BREAK

    # 用一个不会出现在 Markdown 文本里的占位符先 split，再 add_break 替换
    parts = re.split(r"<br\s*/?>", text, flags=re.IGNORECASE)
    para = doc.add_paragraph()
    for i, seg in enumerate(parts):
        if i > 0:
            para.add_run().add_break(WD_BREAK.LINE)
        if seg:
            para.add_run(seg)
