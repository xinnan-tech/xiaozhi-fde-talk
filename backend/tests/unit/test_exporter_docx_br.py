"""Word 导出将 Markdown 中的 `<br>` 渲染为真正的换行。"""
from __future__ import annotations

import io
import re
import zipfile

from docx import Document

from app.services.reports.exporter import export


def _docx_xml(docx_bytes: bytes) -> str:
    """打开 zip，提取 word/document.xml 文本。"""
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        return z.read("word/document.xml").decode("utf-8")


def test_exporter_docx_br_renders_as_real_break():
    """单元格里的 '<br>' 在 Word 里应是真换行（w:br type='textWrapping'）而非字面量。"""
    md = (
        "### 表\n"
        "\n"
        "| a | b |\n"
        "| --- | --- |\n"
        "| x | p<br>q |\n"
    )
    docx_bytes, _mime = export(md, fmt="word", language="en")
    xml = _docx_xml(docx_bytes)

    # 1) 真换行：<w:br .../> 至少出现一次
    assert "<w:br" in xml, f"expected <w:br> in docx, got: {xml[:400]}"
    # 2) 段落里同时含 'p' 和 'q'，说明分到了同一段但被换行切开
    assert "p" in xml and "q" in xml
    # 3) 字面量 '<br>' 不应在 docx 里出现
    assert "&lt;br" not in xml, "found literal '<br>' (escaped) — exporter 仍按字面输出"
    assert ">br<" not in xml


def test_exporter_docx_no_br_no_break():
    """没有 <br> 时不应插任何 <w:br>。"""
    md = "### 标题\n\n普通段落，没有换行符。\n"
    docx_bytes, _mime = export(md, fmt="word", language="en")
    xml = _docx_xml(docx_bytes)
    assert "<w:br" not in xml


def test_exporter_docx_multiple_br_in_one_line():
    """单行内多个 <br> 都应该转成多个换行。"""
    md = "### T\n\n| a | b |\n| --- | --- |\n| x | a<br>b<br>c |\n"
    docx_bytes, _mime = export(md, fmt="word", language="en")
    xml = _docx_xml(docx_bytes)
    # 三个 break：a 与 b 间、b 与 c 间——2 次换行分割出 3 段
    br_count = len(re.findall(r"<w:br", xml))
    assert br_count >= 2, f"expected at least 2 <w:br>, got {br_count}"


def test_exporter_docx_br_in_heading_and_list():
    """标题和列表项中的 `<br>` 也应渲染为真正的换行。"""
    md = "# first<br>second\n\n- item<br>continued\n"
    docx_bytes, _mime = export(md, fmt="word", language="en")
    xml = _docx_xml(docx_bytes)
    assert len(re.findall(r"<w:br", xml)) >= 2
    assert "&lt;br" not in xml


def test_exporter_docx_keeps_escaped_block_prefix_as_text():
    """被 skill 保护的块级前缀在 Word 中应作为普通文本显示。"""
    md = "### T\n\n\\# heading\n\\- item\n1\\. numbered\n"
    docx_bytes, _mime = export(md, fmt="word", language="en")
    paragraphs = [p.text for p in Document(io.BytesIO(docx_bytes)).paragraphs]

    assert "# heading" in paragraphs
    assert "- item" in paragraphs
    assert "1. numbered" in paragraphs


def test_exporter_docx_unescapes_skill_table_cells():
    """Word must show cell values, not the Markdown escapes used in the source."""
    md = (
        "### T\n"
        "\n"
        "| a | b |\n"
        "| --- | --- |\n"
        "| x\\|y | C:\\\\tmp |\n"
    )
    docx_bytes, _mime = export(md, fmt="word", language="en")
    paragraphs = [p.text for p in Document(io.BytesIO(docx_bytes)).paragraphs]

    assert "| x|y | C:\\tmp |" in paragraphs
    assert "| x\\|y | C:\\\\tmp |" not in paragraphs
